"""This script reads from three sources:

 1. a JSON database of key:value data from a protocols.io run
 2. a directory containing files that were once stored in the OSF record relating to the run
 3. a template docx document.

and outputs a single docx document containing data from all three."""

import json
import logging
import re
from typing import Dict, Any

from docx import Document
from docxcompose.composer import Composer

logging.basicConfig()
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

class DocxReplacementTemplateEngine:
    """Processes a docx file to replace entries with [[Filename.docx]] with the contents of
    Filename.docx, reading source files from a directory.

    This is a rudimentary template engine allowing people to concatenate docx files according to
    the content of a master file.  This approach is simple but limited by python's docx processing
    abilities.

    Example usage:

        template_engine = DocxReplacementTemplateEngine("data/word_docs_in_here")
        template_engine.replace("output/generated_doc.docx")
    """

    def __init__(self, source_dir: str, vars_dict: Dict[str, Any]) -> None:
        """Constructor.

        `source_dir`: The directory to load files from for inclusion into the target document
        `vars_dict`: The dictionary of variables to use when replacing values"""
        self.source_dir = source_dir
        self.vars = vars_dict

    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    def replace(self, input_filename: str, output_filename: str) -> None:
        """Process the document at the filename given, replacing placeholders with the contents
        of documents from self.source_dir.

        `document_filename`: The docx filename to process.  This will be written back to the same
                             filename"""

        # Open the source document
        doc = Document(input_filename)
        documents_to_append = []

        for para in doc.paragraphs:
            logger.debug("Processing paragraph with text '%s'", para.text)

            # Replace all expressions
            while (match := re.search("{{(.*?)}}", para.text)) is not None:
                text_before = para.text[:match.start()]
                text_after = para.text[match.end():]
                expression = match.group(1)
                logger.debug("Found expression %s with match %s", expression, match)

                # Find-replace smart quotes
                expression = expression.replace("‘", "'")
                expression = expression.replace("’", "'")
                expression = expression.replace("“", "\"")
                expression = expression.replace("”", "\"")

                # Evaluate expression
                value = ""
                try:
                    value = eval(expression, {}, self.vars)
                except Exception as e:
                    logger.warning(f"WARNING: error evaluating template expression '{expression}': {e}")
                logger.debug("Computed value '%s'", value)

                para.text = f"{text_before}{value}{text_after}"

            # print(f"-> {para.text}")
            if para.text.startswith("[") and para.text.endswith("]"):
                filename = f"{self.source_dir}/{para.text[1:-1].strip()}"
                logger.info(f"Inserting document '{filename}'")
                documents_to_append.append(filename)

                # TODO: delete the paragraph entirely.
                para.text = ""

                # This code attempts to read and re-insert values in-order.
                # is it unreliable, so the decision was made to use the Composer instead.

                # doc_to_insert = Document(f"{self.source_dir}/{filename}")
                # for paragraph_to_insert in doc_to_insert.paragraphs:
                #     inserted_paragraph = doc._body._body._insert_p(paragraph_to_insert._p)
                #     # new_paragraph = para.insert_paragraph_before(paragraph_to_insert.text)
                #     # TODO: delete placeholder, and format the new paragraph
                #     para.text = ""

                # for table_to_insert in doc_to_insert.tables:
                #     print(f"Inserting table")
                #     new_table = doc.add_table(1, 1)
                #     new_table._tbl = table_to_insert._tbl

        composer = Composer(doc)
        for to_append_filename in documents_to_append:
            doc_to_append = Document(to_append_filename)
            composer.append(doc_to_append)
        composer.save(output_filename)



CONFIG_FILENAME = "config.json"

# Load config and data downloaded from previous stages
logger.info(f"Looking for config at {CONFIG_FILENAME}...")
with open(CONFIG_FILENAME) as fin:
    config = json.load(fin)

with open(config["database_filename"]) as fin:
    db = json.load(fin)

logger.info(f"Using template at {config['template_filename']}")
template_engine = DocxReplacementTemplateEngine(config['docx_files_dirname'], db)
template_engine.replace(config['template_filename'], config['output_filename'])
logger.info(f"Success.  Output can be found at {config['output_filename']}")
